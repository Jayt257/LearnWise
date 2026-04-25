"""
LearnWise — Software Design Document Generator
Generates DESIGN_DOCUMENT.docx for stakeholder presentation.
Run: python docs/generate_design_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)


# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Set table cell background fill colour (hex string e.g. '1F3864')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_heading(text, level=1, color="1F3864"):
    """Add a styled heading paragraph."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(
            int(color[0:2], 16),
            int(color[2:4], 16),
            int(color[4:6], 16)
        )
        if level == 1:
            run.font.size = Pt(18)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.bold = True
    return p


def add_body(text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code_block(text):
    """Add a monospace code-style paragraph with shaded background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    # Light grey shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    return p


def add_table(headers, rows, header_color="1F3864", header_text_color="FFFFFF"):
    """Add a formatted table."""
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_color)
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.size  = Pt(9)
                run.font.color.rgb = RGBColor(
                    int(header_text_color[0:2], 16),
                    int(header_text_color[2:4], 16),
                    int(header_text_color[4:6], 16)
                )

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg = "FFFFFF" if r_idx % 2 == 0 else "EBF2FA"
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val)
            set_cell_bg(cells[c_idx], bg)
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    return table


def add_note(text, prefix="NOTE: "):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(prefix)
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x19, 0x63, 0x84)
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    r2.italic = True
    return p


def page_break():
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("LearnWise")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("AI-Powered Language Learning Platform")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x27, 0x6E, 0xD8)

doc.add_paragraph()

title2_p = doc.add_paragraph()
title2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2_p.add_run("SOFTWARE DESIGN DOCUMENT")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph()
doc.add_paragraph()

meta_entries = [
    ("Document Version",    "1.0"),
    ("Standard",            "IEEE Std 1016 — Software Design Descriptions"),
    ("Classification",      "Technical Design Reference — Stakeholder Presentation"),
    ("Prepared For",        "Manager & Client Review"),
    ("Backend Stack",       "FastAPI (Python 3.10) · PostgreSQL · SQLAlchemy ORM"),
    ("AI/ML Stack",         "Groq LLaMA-3-8B-8192 · OpenAI Whisper (on-premise)"),
    ("Frontend Stack",      "React 18 · Redux Toolkit · Axios · Vite"),
    ("Date",                "April 2026"),
]

meta_table = doc.add_table(rows=len(meta_entries), cols=2)
meta_table.style = 'Table Grid'
for i, (k, v) in enumerate(meta_entries):
    cells = meta_table.rows[i].cells
    cells[0].text = k
    cells[1].text = v
    bg = "1F3864" if i == 0 else ("EBF2FA" if i % 2 == 0 else "FFFFFF")
    txt_color = "FFFFFF" if i == 0 else "000000"
    set_cell_bg(cells[0], bg)
    set_cell_bg(cells[1], bg)
    for c in cells:
        for para in c.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.bold = (i == 0)
                run.font.color.rgb = RGBColor(
                    int(txt_color[0:2], 16),
                    int(txt_color[2:4], 16),
                    int(txt_color[4:6], 16)
                )

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS (manual)
# ─────────────────────────────────────────────────────────────────────────────
add_heading("Table of Contents", level=1)

toc_entries = [
    ("1.", "System Context & Architecture Overview"),
    ("2.", "Section 1: Class Diagrams"),
    ("   2.1", "Backend Database Models Layer"),
    ("   2.2", "Backend Core & Services Layer"),
    ("   2.3", "Backend API Routers Layer"),
    ("   2.4", "Frontend Redux Store & API Client Layer"),
    ("   2.5", "Frontend Pages & Components Layer"),
    ("3.", "Section 2: Sequence Diagrams"),
    ("   3.1 – 3.11", "All 11 Sequence Flows (Registration → Server Startup)"),
    ("4.", "Section 3: Activity Diagrams"),
    ("   4.1 – 4.8", "All 8 Activity Flows (Registration → Admin CRUD)"),
    ("5.", "Section 4: State Diagrams"),
    ("   5.1 – 5.13", "All 13 State Machines (User Lifecycle → AdminDashboard)"),
    ("6.", "Complete Entity Reference Table"),
]

for num, text in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{num}  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 0: SYSTEM CONTEXT
# ─────────────────────────────────────────────────────────────────────────────
add_heading("1. System Context & Architecture Overview", level=1)

add_body(
    "LearnWise is a full-stack AI-powered language learning platform. Users progress through a "
    "hierarchical curriculum: Language Pair → Month → Block → Activity. Each block contains 8 "
    "distinct activity types graded by either local MCQ scoring or the Groq LLaMA-3 AI API. "
    "Audio-based activities (Speaking, Pronunciation) are transcribed using an on-premise OpenAI "
    "Whisper model. All progression data is persisted to PostgreSQL, with XP accumulation and "
    "social leaderboards driving user engagement."
)

doc.add_paragraph()
add_body("Architecture Layers", bold=True)
add_code_block(
    "Browser (React 18 + Redux Toolkit + Axios)\n"
    "         ↕  HTTPS / REST JSON\n"
    "FastAPI Backend (Uvicorn · Python 3.10)\n"
    "   ├── Routers: Auth · Users · Progress · Validate · Speech · Leaderboard · Friends · Admin\n"
    "   ├── Services: GroqService · WhisperService · ScoringService · ContentService\n"
    "   └── Core: Security (JWT+bcrypt) · Database (SQLAlchemy ORM) · Config · Dependencies\n"
    "         ↕  SQLAlchemy\n"
    "PostgreSQL Database\n"
    "         ↕  External HTTP\n"
    "Groq Cloud API (LLaMA-3-8B)       OpenAI Whisper (on-premise, local model)"
)

add_body("Curriculum Hierarchy", bold=True)
add_code_block(
    "Language Pair (e.g. hi-ja  Hindu → Japanese)\n"
    "  └── Month  (1 – 3)\n"
    "        └── Block  (1 – 6 per month)\n"
    "              └── Activity  (8 per block: lesson, vocabulary, pronunciation,\n"
    "                                          reading, writing, listening, speaking, test)\n"
    "Total activities per pair: 3 × 6 × 8 = 144  (activity_seq_id: 1 – 144)"
)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1: CLASS DIAGRAMS
# ─────────────────────────────────────────────────────────────────────────────
add_heading("2. Section 1: Class Diagrams", level=1)
add_note(
    "Class diagrams follow UML 2.5 notation. Stereotypes indicate technology layers. "
    "Each class lists all attributes (with types and constraints), all methods, and all "
    "inter-class associations with multiplicities.",
    prefix="Notation: "
)
doc.add_paragraph()

# ── 2.1 DB Models ──────────────────────────────────────────────────────────
add_heading("2.1  Backend Database Models Layer", level=2)
add_body("Files: backend/app/models/user.py  ·  backend/app/models/progress.py  ·  backend/app/models/friends.py", italic=True)
doc.add_paragraph()

# User class table
add_body("Class: User   «SQLAlchemy ORM»   Table: users", bold=True)
add_table(
    ["Attribute", "Type", "Constraint"],
    [
        ["id",             "UUID (PK)",           "auto-generated, uuid4"],
        ["username",       "String(50)",          "UNIQUE, NOT NULL, indexed"],
        ["email",          "String(255)",         "UNIQUE, NOT NULL, indexed"],
        ["password_hash",  "String(255)",         "NOT NULL"],
        ["display_name",   "String(100)",         "nullable"],
        ["avatar_url",     "String(500)",         "nullable"],
        ["native_lang",    "String(5)",           "nullable, default 'en'"],
        ["role",           "Enum(UserRole)",      "NOT NULL, default 'user'"],
        ["is_active",      "Boolean",             "default True"],
        ["created_at",     "DateTime",            "default utcnow"],
        ["last_active",    "DateTime",            "default utcnow, on update utcnow"],
    ]
)
add_body("Methods: __repr__() → str", italic=True)
add_body("Relationships:", bold=True)
add_code_block(
    "User  \"1\" ───────────── \"*\"  UserLanguageProgress   (language_progress, cascade delete)\n"
    "User  \"1\" ───────────── \"*\"  ActivityCompletion     (activity_completions, cascade delete)\n"
    "User  \"1\" ───────────── \"*\"  FriendRequest          (sent_requests, cascade delete)\n"
    "User  \"1\" ───────────── \"*\"  FriendRequest          (received_requests, cascade delete)"
)

add_body("Enum: UserRole", bold=True)
add_table(["Value"], [["user"], ["admin"]])

doc.add_paragraph()

# UserLanguageProgress
add_body("Class: UserLanguageProgress   «SQLAlchemy ORM»   Table: user_language_progress", bold=True)
add_note("Unique Constraint: (user_id, lang_pair_id)")
add_table(
    ["Attribute", "Type", "Constraint"],
    [
        ["id",                  "UUID (PK)",          "auto-generated"],
        ["user_id",             "UUID (FK→users.id)", "NOT NULL, CASCADE DELETE"],
        ["lang_pair_id",        "String(10)",         "NOT NULL  e.g. 'hi-ja'"],
        ["total_xp",            "Integer",            "default 0"],
        ["current_month",       "Integer",            "default 1"],
        ["current_block",       "Integer",            "default 1"],
        ["current_activity_id", "Integer",            "default 1, global seq 1–144"],
        ["started_at",          "DateTime",           "default utcnow"],
        ["last_activity_at",    "DateTime",           "default utcnow, on update"],
    ]
)

add_body("Class: ActivityCompletion   «SQLAlchemy ORM»   Table: activity_completions", bold=True)
add_note("Unique Constraint: (user_id, lang_pair_id, activity_seq_id)")
add_table(
    ["Attribute", "Type", "Constraint"],
    [
        ["id",               "UUID (PK)",          "auto-generated"],
        ["user_id",          "UUID (FK→users.id)", "NOT NULL, CASCADE DELETE"],
        ["lang_pair_id",     "String(10)",         "NOT NULL"],
        ["activity_seq_id",  "Integer",            "NOT NULL, default 1"],
        ["activity_json_id", "String(80)",         "nullable  e.g. 'ja_hi_M1B1_lesson_1'"],
        ["activity_type",    "String(20)",         "NOT NULL: lesson/vocabulary/reading/writing/listening/speaking/pronunciation/test"],
        ["month_number",     "Integer",            "nullable"],
        ["block_number",     "Integer",            "nullable"],
        ["score_earned",     "Integer",            "default 0"],
        ["max_score",        "Integer",            "NOT NULL"],
        ["passed",           "Boolean",            "default False"],
        ["attempts",         "Integer",            "default 1"],
        ["ai_feedback",      "Text",               "nullable"],
        ["ai_suggestion",    "Text",               "nullable"],
        ["completed_at",     "DateTime",           "default utcnow"],
    ]
)

add_body("Class: FriendRequest   «SQLAlchemy ORM»   Table: friend_requests", bold=True)
add_note("Unique Constraint: (sender_id, receiver_id)")
add_table(
    ["Attribute", "Type", "Constraint"],
    [
        ["id",          "UUID (PK)",              "auto-generated"],
        ["sender_id",   "UUID (FK→users.id)",    "NOT NULL, CASCADE DELETE"],
        ["receiver_id", "UUID (FK→users.id)",    "NOT NULL, CASCADE DELETE"],
        ["status",      "Enum(FriendRequestStatus)", "NOT NULL, default 'pending'"],
        ["created_at",  "DateTime",              "default utcnow"],
        ["updated_at",  "DateTime",              "default utcnow, on update"],
    ]
)
add_body("Enum: FriendRequestStatus  —  values: pending | accepted | declined", italic=True)

page_break()

# ── 2.2 Core & Services ────────────────────────────────────────────────────
add_heading("2.2  Backend Core & Services Layer", level=2)
add_body("Files: backend/app/core/  ·  backend/app/services/", italic=True)
doc.add_paragraph()

services = [
    {
        "name": "Settings   «Pydantic BaseSettings»",
        "file": "backend/app/core/config.py",
        "attrs": [
            ["APP_NAME",                     "str",  "default 'LearnWise'"],
            ["DEBUG",                         "bool", "default True"],
            ["DATABASE_URL",                  "str",  "required (from .env)"],
            ["SECRET_KEY",                    "str",  "required"],
            ["ALGORITHM",                     "str",  "default 'HS256'"],
            ["ACCESS_TOKEN_EXPIRE_MINUTES",   "int",  "default 10080 (7 days)"],
            ["GROQ_API_KEY",                  "str",  "required"],
            ["GROQ_MODEL",                    "str",  "default 'llama3-8b-8192'"],
            ["WHISPER_MODEL",                 "str",  "default 'small'"],
            ["ALLOWED_ORIGINS",               "str",  "default 'http://localhost:5173'"],
            ["ADMIN_EMAIL",                   "str",  "default 'admin@learnwise.app'"],
            ["ADMIN_PASSWORD",                "str",  "default 'Admin@LearnWise2026'"],
            ["ADMIN_USERNAME",                "str",  "default 'admin'"],
            ["SCORE_THRESHOLD_OVERRIDE",      "int",  "default -1 (disabled); set to 0 for test suites"],
            ["DATA_DIR",                      "str",  "default 'data'"],
        ],
        "methods": ["origins_list() → List[str]", "data_path() → str"],
    },
    {
        "name": "Security   «Module»",
        "file": "backend/app/core/security.py",
        "attrs": [["pwd_context", "CryptContext", "bcrypt scheme"]],
        "methods": [
            "hash_password(password: str) → str",
            "verify_password(plain: str, hashed: str) → bool",
            "create_access_token(data: dict, expires_delta: Optional[timedelta]) → str",
            "decode_token(token: str) → Optional[dict]  — returns payload or None on JWTError",
        ],
    },
    {
        "name": "Dependencies   «Module»",
        "file": "backend/app/core/dependencies.py",
        "attrs": [],
        "methods": [
            "get_current_user(credentials, db) → User  — decodes Bearer JWT, returns active User or 401",
            "get_current_active_user(user) → User",
            "require_admin(user) → User  — checks role == 'admin' or raises 403",
        ],
    },
    {
        "name": "ScoringService   «Service Module»",
        "file": "backend/app/services/scoring_service.py",
        "attrs": [["PASS_THRESHOLD", "float", "0.2  (20% of max_xp)"]],
        "methods": [
            "calculate_score(questions, max_xp, groq_scores) → dict  — clamps per-question scores to [0, per_q_max], sums, computes percentage",
            "score_mcq_locally(questions, max_xp) → dict  — exact string match: str(user_answer)==str(correct_answer)",
        ],
    },
    {
        "name": "GroqService   «Service Module»",
        "file": "backend/app/services/groq_service.py",
        "attrs": [
            ["RUBRICS",     "dict", "per activity_type rubric strings"],
            ["TIER_CONFIG", "dict", "feedback_tier → prompt instructions (hint/lesson/praise)"],
        ],
        "methods": [
            "get_client() → Groq",
            "_determine_feedback_tier(pct, attempt_count) → str — hint/lesson/praise",
            "_build_prompt(activity_type, questions, max_xp, user_lang, target_lang, tier) → str",
            "validate_activity(type, questions, max_xp, user_lang, target_lang, attempt_count) → dict",
            "generate_tier_feedback(type, tier, overall_feedback, suggestion, ...) → dict",
        ],
    },
    {
        "name": "WhisperService   «Service Module»",
        "file": "backend/app/services/whisper_service.py",
        "attrs": [
            ["_whisper_model",     "Any",  "cached model reference (None initially)"],
            ["_whisper_available", "bool", "cached availability flag (None initially)"],
        ],
        "methods": [
            "_load_model() → bool  — lazy-loads Whisper model once",
            "transcribe_audio(audio_bytes, filename) → dict  — returns {text, language, confidence, is_mock}",
            "save_audio_file(audio_bytes, filename, pair_id) → str  — saves WAV to uploads/{pair_id}/",
        ],
    },
    {
        "name": "ContentService   «Service Module»",
        "file": "backend/app/services/content_service.py",
        "attrs": [],
        "methods": [
            "get_all_pairs() → List[dict]          get_meta(pair_id) → dict",
            "get_content_file(pair_id, file_path) → dict",
            "scaffold_pair(pair_id)                — creates full 3×6×8=144 JSON structure",
            "add_month(pair_id, month_num)          add_block(pair_id, month_num, block_num)",
            "delete_pair(pair_id)                   delete_month(pair_id, month_num)",
            "delete_block(pair_id, month_num, block_num)",
            "get_analytics() → dict                — total pairs, activities per type",
        ],
    },
]

for svc in services:
    add_body(f"Class: {svc['name']}", bold=True)
    add_body(f"File: {svc['file']}", italic=True)
    if svc["attrs"]:
        add_table(["Attribute / Constant", "Type", "Notes"], svc["attrs"])
    if svc["methods"]:
        add_body("Methods:", bold=True)
        for m in svc["methods"]:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(m)
            run.font.size = Pt(9)
            run.font.name = 'Courier New'
    doc.add_paragraph()

page_break()

# ── 2.3 Routers ────────────────────────────────────────────────────────────
add_heading("2.3  Backend API Routers Layer", level=2)
add_body("Files: backend/app/routers/", italic=True)
doc.add_paragraph()

routers = [
    ("AuthRouter", "/api/auth", "auth.py", [
        "POST /register   → register(req: RegisterRequest) → TokenResponse",
        "POST /login      → login(req: LoginRequest) → TokenResponse",
        "POST /admin/login → admin_login(req: AdminLoginRequest) → TokenResponse",
        "GET  /me         → get_me(current_user) → UserOut",
    ]),
    ("UsersRouter", "/api/users", "users.py", [
        "GET  /me           → get_my_profile() → UserProfileOut",
        "PUT  /me           → update_profile(req) → UserProfileOut",
        "GET  /search       → search_users(q, limit) → UserSearchResult",
        "GET  /{user_id}    → get_user_profile(user_id) → UserPublicOut",
        "GET  /{user_id}/progress → get_user_progress_public(user_id) → dict",
    ]),
    ("ProgressRouter", "/api/progress", "progress.py", [
        "GET  /                        → get_all_progress() → List[ProgressOut]",
        "GET  /{pair_id}               → get_pair_progress(pair_id) → ProgressOut",
        "POST /{pair_id}/start         → start_pair(pair_id) → ProgressOut",
        "POST /{pair_id}/complete      → complete_activity(pair_id, req) → CompletionOut",
        "GET  /{pair_id}/completions   → get_completions(pair_id) → List[CompletionOut]",
        "PRIVATE: _derive_month_block(activity_seq_id) → tuple[int, int]",
    ]),
    ("ValidateRouter", "/api/validate", "validate.py", [
        "POST /  → validate_activity(req: ValidateRequest) → ValidateResponse",
        "CONSTANTS: MCQ_TYPES = {'test'}",
        "CONSTANTS: GROQ_TYPES = {'writing','speaking','pronunciation','listening','reading','vocabulary','lesson'}",
    ]),
    ("SpeechRouter", "/api/speech", "speech.py", [
        "POST /transcribe → transcribe(audio: UploadFile) → TranscribeResponse",
        "CONSTANTS: MAX_AUDIO_SIZE=25MB, MIN_AUDIO_SIZE=100 bytes",
        "PRIVATE: _is_allowed_audio(content_type) → bool",
    ]),
    ("LeaderboardRouter", "/api/leaderboard", "leaderboard.py", [
        "GET /{pair_id}         → get_leaderboard(pair_id, limit=50) → List[LeaderboardEntry]",
        "GET /{pair_id}/friends → get_friends_leaderboard(pair_id) → List[LeaderboardEntry]",
    ]),
    ("FriendsRouter", "/api/friends", "friends.py", [
        "GET  /                          → get_friends() → dict",
        "GET  /requests                  → get_incoming_requests() → dict",
        "POST /request/{user_id}         → send_request(user_id) → dict",
        "PUT  /request/{req_id}/accept   → accept_request(req_id) → dict",
        "PUT  /request/{req_id}/decline  → decline_request(req_id) → dict",
        "DELETE /{user_id}               → remove_friend(user_id) → dict",
    ]),
    ("AdminRouter", "/api/admin", "admin.py  [requires require_admin dependency]", [
        "GET  /users                                          → get_all_users() → List[UserOut]",
        "POST /users/{id}/activate                            → activate_user()",
        "POST /users/{id}/deactivate                          → deactivate_user()",
        "GET  /analytics                                      → get_analytics() → dict",
        "GET  /content/pairs                                  → get_pairs() → List[dict]",
        "POST /content/pairs                                  → create_pair(pair_id) → dict",
        "DELETE /content/pairs/{pair_id}                     → delete_pair() → dict",
        "POST /content/pairs/{pair_id}/months                → add_month()",
        "DELETE /content/pairs/{pair_id}/months/{month_num}  → delete_month()",
        "POST .../blocks                                      → add_block()",
        "DELETE .../blocks/{block_num}                       → delete_block()",
        "GET/PUT/POST/DELETE .../activities/{type}           → CRUD on activity JSON files",
        "GET  /stats                                          → get_admin_stats() → dict",
    ]),
]

for name, prefix, file_, endpoints in routers:
    add_body(f"Class: {name}   «FastAPI Router»", bold=True)
    add_body(f"URL Prefix: {prefix}   |   File: {file_}", italic=True)
    for ep in endpoints:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(ep)
        run.font.size = Pt(9)
        run.font.name = 'Courier New'
    doc.add_paragraph()

page_break()

# ── 2.4 Redux & API Clients ────────────────────────────────────────────────
add_heading("2.4  Frontend Redux Store & API Client Layer", level=2)
add_body("Files: src/store/  ·  src/api/", italic=True)
doc.add_paragraph()

slices = [
    {
        "name": "AuthSlice   «Redux Slice»",
        "file": "src/store/authSlice.js",
        "state": [
            ["token",           "string | null",  "JWT or null"],
            ["user",            "UserOut | null",  "current user object"],
            ["isAuthenticated", "boolean",         "login state"],
            ["loading",         "boolean",         "API request in progress"],
            ["error",           "string | null",   "last error message"],
        ],
        "actions": ["logout()", "updateUser(payload)", "clearError()"],
        "thunks": ["loginUser(credentials)", "registerUser(data)", "adminLoginUser(credentials)"],
    },
    {
        "name": "ProgressSlice   «Redux Slice»",
        "file": "src/store/progressSlice.js",
        "state": [
            ["allProgress",    "Array[ProgressOut]", "all language pair progress"],
            ["pairs",          "Array[{id,...}]",     "available language pairs"],
            ["currentPairId",  "string | null",       "active selected pair"],
            ["loading",        "boolean",             ""],
            ["error",          "string | null",       ""],
        ],
        "actions": ["setCurrentPair(pairId)", "updateProgressXP({pairId, xpDelta})", "advanceProgress({pairId, newProgress})"],
        "thunks": ["fetchAllProgress()", "fetchPairProgress(pairId)", "startLanguagePair(pairId)", "fetchPairs()"],
    },
]

for sl in slices:
    add_body(f"Class: {sl['name']}", bold=True)
    add_body(f"File: {sl['file']}", italic=True)
    add_body("State Shape:", bold=True)
    add_table(["Field", "Type", "Notes"], sl["state"])
    add_body("Sync Reducers:", bold=True)
    for a in sl["actions"]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(a)
        run.font.size = Pt(9)
    add_body("Async Thunks:", bold=True)
    for t in sl["thunks"]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(t)
        run.font.size = Pt(9)
    doc.add_paragraph()

add_body("Class: ApiClient   «Axios Instance»   File: src/api/client.js", bold=True)
add_table(
    ["Property", "Value / Behaviour"],
    [
        ["baseURL",              "${VITE_API_BASE_URL}/api"],
        ["Content-Type",         "application/json"],
        ["Request Interceptor",  "Attaches Authorization: Bearer {lw_token} from localStorage"],
        ["Response Interceptor", "On 401: clears localStorage → redirects admin to /admin/login, user to /login"],
    ]
)

api_modules = [
    ("AuthApiModule",      "src/api/auth.js",     ["login(credentials)", "register(data)", "adminLogin(credentials)"]),
    ("ProgressApiModule",  "src/api/progress.js", ["getAllProgress()", "getPairProgress(pairId)", "startPair(pairId)", "completeActivity(pairId, payload)", "validateActivity(payload)", "getCompletions(pairId)"]),
    ("ContentApiModule",   "src/api/content.js",  ["getPairs()", "getMeta(pairId)"]),
    ("UsersApiModule",     "src/api/users.js",    ["getMyProfile()", "updateProfile(data)", "searchUsers(q)", "getUserProfile(userId)", "getLeaderboard(pairId)", "getFriends()", "sendFriendRequest(userId)"]),
]

for name, file_, methods in api_modules:
    add_body(f"Class: {name}   «API Module»   File: {file_}", bold=True)
    for m in methods:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(m)
        run.font.size = Pt(9)
        run.font.name = 'Courier New'
    doc.add_paragraph()

page_break()

# ── 2.5 Frontend Pages ─────────────────────────────────────────────────────
add_heading("2.5  Frontend Pages & Components Layer", level=2)
doc.add_paragraph()

add_body("Page Components", bold=True)
add_table(
    ["Component", "File", "Key State / Logic"],
    [
        ["App",               "src/App.jsx",          "Routes: /login, /register, /admin/login, /onboarding, /dashboard, /activity/:pairId/:type, /leaderboard, /profile, /search, /admin"],
        ["DashboardPage",     "DashboardPage.jsx",    "State: meta, completions, selectedPair, loading. Logic: isUnlocked(seqId), isCompleted(seqId)"],
        ["ActivityRouter",    "ActivityRouter.jsx",   "Props from navigate state: activityFile, activitySeqId, maxXP, monthNumber, blockNumber. Routes to 8 page components"],
        ["LessonPage",        "LessonPage.jsx",       "Renders lesson content sections + comprehension questions"],
        ["VocabPage",         "VocabPage.jsx",        "Flashcard + MCQ vocabulary activity"],
        ["ReadingPage",       "ReadingPage.jsx",      "Passage display + short answer fields"],
        ["WritingPage",       "WritingPage.jsx",      "Free-text essay submission"],
        ["ListeningPage",     "ListeningPage.jsx",    "AudioPlayer component + MCQ answers"],
        ["SpeakingPage",      "SpeakingPage.jsx",     "AudioRecorder component + Whisper STT"],
        ["PronunciationPage", "PronunciationPage.jsx","STT pipeline + phonetic comparison"],
        ["TestPage",          "TestPage.jsx",         "Multi-type MCQ assessment"],
        ["LoginPage",         "src/pages/auth/",      "Dispatches: loginUser()"],
        ["RegisterPage",      "src/pages/auth/",      "Dispatches: registerUser()"],
        ["AdminLoginPage",    "src/pages/admin/",     "Dispatches: adminLoginUser()"],
        ["LeaderboardPage",   "src/pages/social/",   "Fetches: /api/leaderboard/{pair} + /friends"],
        ["ProfilePage",       "src/pages/social/",   "Fetches: /api/users/me, /api/progress, completions"],
        ["SearchFriendsPage", "src/pages/social/",   "Fetches: /api/users/search, friend requests"],
        ["OnboardingPage",    "src/pages/onboarding/","Dispatches: startLanguagePair(pairId)"],
        ["AdminDashboard",    "AdminDashboard.jsx",   "All Admin API endpoints: user management + curriculum CRUD"],
    ]
)

add_body("Shared UI Components", bold=True)
add_table(
    ["Component", "File", "Purpose"],
    [
        ["Sidebar",          "Sidebar.jsx",          "Navigation links + logout action"],
        ["ProtectedRoute",   "ProtectedRoute.jsx",   "Redirects unauthenticated users to /login"],
        ["AdminRoute",       "AdminRoute.jsx",        "Redirects non-admins to /admin/login"],
        ["AudioPlayer",      "AudioPlayer.jsx",       "Renders <audio> from {API_URL}/static/{path}; hides when path=null"],
        ["AudioRecorder",    "AudioRecorder.jsx",     "Captures mic audio via MediaRecorder → POSTs to /api/speech/transcribe"],
        ["DynamicQuiz",      "DynamicQuiz.jsx",       "Renders MCQ / fill-blank / short-answer blocks from activity JSON"],
        ["ScoreModal",       "ScoreModal.jsx",        "Shows pass/fail result, XP earned, AI feedback, suggestion"],
        ["ActivityFeedback", "ActivityFeedback.jsx",  "Feedback + suggestion display component"],
        ["XPBurst",          "XPBurst.jsx",           "Animated XP gained notification overlay"],
        ["ErrorBoundary",    "ErrorBoundary.jsx",     "React error boundary wrapper for uncaught render errors"],
        ["NebulaWeb",        "NebulaWeb.jsx",         "Animated background canvas effect (glassmorphism theme)"],
        ["ThemeToggle",      "ThemeToggle.jsx",       "Light / dark mode switcher"],
    ]
)

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2: SEQUENCE DIAGRAMS
# ─────────────────────────────────────────────────────────────────────────────
add_heading("3. Section 2: Sequence Diagrams", level=1)
add_note(
    "Sequence diagrams document time-ordered message exchanges between actors and system components. "
    "ALT frames show conditional branches. PAR frames show parallel operations. "
    "Return arrows (→) show synchronous responses.",
    prefix="Notation: "
)
doc.add_paragraph()

sequences = [
    {
        "id": "SD-01",
        "title": "User Registration",
        "trigger": "User fills Register form and clicks Submit",
        "route": "POST /api/auth/register",
        "participants": ["User", "React Frontend", "Redux Store", "Axios Client", "Auth Router", "Security Module", "PostgreSQL DB"],
        "steps": [
            ("User", "React Frontend",   "fills {username, email, password}, clicks Register"),
            ("React Frontend", "Redux Store", "dispatch(registerUser(credentials))"),
            ("Redux Store",    "Axios Client",    "POST /api/auth/register payload"),
            ("Axios Client",   "Auth Router",     "HTTP POST /api/auth/register"),
            ("Auth Router",    "PostgreSQL DB",   "query User WHERE email == req.email"),
            ("[ALT]", "",                "IF email exists → return 400 'Email already registered'"),
            ("[ALT]", "",                "IF username taken → return 400 'Username already taken'"),
            ("Auth Router",    "Security Module", "hash_password(req.password)"),
            ("Security Module","Auth Router",     "bcrypt hash string"),
            ("Auth Router",    "PostgreSQL DB",   "INSERT User {username, email, hash, role='user'}"),
            ("Auth Router",    "Security Module", "create_access_token({sub: user.id, role: 'user'})"),
            ("Security Module","Auth Router",     "HS256 JWT (7-day expiry)"),
            ("Auth Router",    "Axios Client",    "201 TokenResponse {access_token, token_type, user}"),
            ("Axios Client",   "Redux Store",     "resolved payload"),
            ("Redux Store",    "localStorage",    "setItem('lw_token'), setItem('lw_user')"),
            ("Redux Store",    "React Frontend",  "state.isAuthenticated = true"),
            ("React Frontend", "User",            "navigate('/dashboard')"),
        ],
    },
    {
        "id": "SD-02",
        "title": "User Login",
        "trigger": "User fills Login form and clicks Login",
        "route": "POST /api/auth/login",
        "participants": ["User", "React Frontend", "Axios Client", "Auth Router", "Security Module", "PostgreSQL DB"],
        "steps": [
            ("User",           "React Frontend",  "fills {email, password}, clicks Login"),
            ("React Frontend", "Axios Client",    "POST /api/auth/login"),
            ("Axios Client",   "Auth Router",     "HTTP POST /api/auth/login"),
            ("Auth Router",    "PostgreSQL DB",   "query User WHERE email==req.email AND role=='user'"),
            ("[ALT]", "",                          "User not found → 401 'Invalid email or password'"),
            ("[ALT]", "",                          "user.is_active==False → 403 'Account is deactivated'"),
            ("Auth Router",    "Security Module", "verify_password(req.password, hash) → True"),
            ("Auth Router",    "PostgreSQL DB",   "UPDATE user.last_active = utcnow()"),
            ("Auth Router",    "Security Module", "create_access_token({sub: user.id})"),
            ("Auth Router",    "Axios Client",    "200 TokenResponse"),
            ("Axios Client",   "React Frontend",  "isAuthenticated = true"),
            ("React Frontend", "User",            "navigate('/dashboard')"),
        ],
    },
    {
        "id": "SD-03",
        "title": "JWT Authentication on Every Protected Request",
        "trigger": "Any API call to a protected endpoint",
        "route": "ALL /api/* except /auth/register and /auth/login",
        "participants": ["Axios Client", "FastAPI Backend", "Dependencies", "Security Module", "PostgreSQL DB", "localStorage"],
        "steps": [
            ("Axios Client",    "localStorage",   "getItem('lw_token')"),
            ("Axios Client",    "FastAPI Backend", "HTTP Request + Authorization: Bearer {token}"),
            ("FastAPI Backend", "Dependencies",   "get_current_user(credentials, db)"),
            ("Dependencies",    "Security Module","decode_token(token)"),
            ("[ALT]", "",                          "Token invalid/expired → JWTError → 401 → Axios removes localStorage entries → redirect to /login or /admin/login"),
            ("Security Module", "Dependencies",   "payload {sub: user_id, role, exp}"),
            ("Dependencies",    "PostgreSQL DB",  "query User WHERE id==user_id AND is_active==True"),
            ("PostgreSQL DB",   "Dependencies",   "User object"),
            ("Dependencies",    "FastAPI Backend","return User (injected into endpoint handler)"),
        ],
    },
    {
        "id": "SD-04",
        "title": "Dashboard Load (Parallel Fetches)",
        "trigger": "User navigates to /dashboard",
        "route": "GET /api/progress, GET /api/content/pairs, GET /api/content/{pair}/meta, GET /api/progress/{pair}/completions",
        "participants": ["User", "DashboardPage", "Redux Store", "Axios Client", "FastAPI Backend", "File System", "PostgreSQL DB"],
        "steps": [
            ("User",              "DashboardPage", "navigate('/dashboard')"),
            ("[PAR]", "",          "dispatch(fetchAllProgress()) || dispatch(fetchPairs())"),
            ("Redux Store",        "FastAPI Backend", "GET /api/progress → List[ProgressOut]"),
            ("Redux Store",        "FastAPI Backend", "GET /api/content/pairs → List[pair IDs]"),
            ("DashboardPage",      "FastAPI Backend", "GET /api/content/{pairId}/meta → meta.json tree"),
            ("DashboardPage",      "FastAPI Backend", "GET /api/progress/{pairId}/completions → List[CompletionOut]"),
            ("DashboardPage",      "DashboardPage",  "Build completedSeqIds = Set(passed completions)"),
            ("DashboardPage",      "DashboardPage",  "isUnlocked(id) = id <= current_activity_id"),
            ("DashboardPage",      "User",            "Render roadmap: colored/locked/completed nodes"),
        ],
    },
    {
        "id": "SD-05",
        "title": "Activity Submission & AI Grading",
        "trigger": "User clicks Submit on an open-ended or MCQ activity",
        "route": "POST /api/validate",
        "participants": ["User", "Activity Page", "Validate Router", "Groq Service", "Scoring Service", "Groq Cloud API"],
        "steps": [
            ("User",            "Activity Page",   "click Submit (all answers filled)"),
            ("Activity Page",   "Validate Router", "POST /api/validate {activity_type, questions, max_xp, ...}"),
            ("[ALT MCQ]", "",                       "activity_type in MCQ_TYPES ('test') → score_mcq_locally: str comparison per question"),
            ("[ALT GROQ]", "",                      "activity_type in GROQ_TYPES → _build_prompt with rubrics"),
            ("Groq Service",    "Groq Cloud API",  "POST completions {llama3-8b-8192, temp=0.1}"),
            ("Groq Cloud API",  "Groq Service",    "JSON {question_results, overall_feedback, suggestion}"),
            ("Groq Service",    "Validate Router", "parsed + merged groq_result"),
            ("Validate Router", "Scoring Service", "calculate_score(questions, max_xp, groq_scores)"),
            ("Validate Router", "Groq Service",    "_determine_feedback_tier(pct, attempt_count)"),
            ("[OPT]", "",                           "tier != 'lesson' → generate_tier_feedback() (second Groq call)"),
            ("[OPT]", "",                           "SCORE_THRESHOLD_OVERRIDE > 0 → effective_passed = score >= OVERRIDE"),
            ("Validate Router", "Activity Page",   "200 ValidateResponse {score, passed, feedback, tier}"),
            ("Activity Page",   "User",            "show ScoreModal with result"),
        ],
    },
    {
        "id": "SD-06",
        "title": "Activity Completion Recording & Progression Unlock",
        "trigger": "After scoring is received, activity page records completion",
        "route": "POST /api/progress/{pairId}/complete",
        "participants": ["Activity Page", "Axios Client", "Progress Router", "PostgreSQL DB", "Redux Store"],
        "steps": [
            ("Activity Page",   "Progress Router", "POST /api/progress/{pairId}/complete {activity_seq_id, score_earned, passed, ...}"),
            ("Progress Router", "Progress Router", "apply SCORE_THRESHOLD_OVERRIDE if > 0"),
            ("Progress Router", "Progress Router", "_derive_month_block(seq_id): month=(zero//48)+1, block=((zero%48)//8)+1"),
            ("Progress Router", "PostgreSQL DB",   "query UserLanguageProgress WHERE user_id AND pair_id"),
            ("[ALT]", "",                           "No record → INSERT new UserLanguageProgress (m=1, b=1, a=seq_id)"),
            ("Progress Router", "PostgreSQL DB",   "query ActivityCompletion WHERE user_id AND pair_id AND seq_id"),
            ("[ALT retry]", "", "Completion exists → new_score > old_score → xp_delta = new-old, UPDATE score"),
            ("[ALT retry]", "", "new_score <= old_score → xp_delta = 0"),
            ("[ALT first]", "", "First attempt → INSERT ActivityCompletion, xp_delta = score_earned"),
            ("Progress Router", "PostgreSQL DB",   "UPDATE total_xp += xp_delta"),
            ("[OPT]", "",                           "passed=True AND seq_id==current_activity_id → current_activity_id++ (UNLOCK next)"),
            ("Progress Router", "PostgreSQL DB",   "COMMIT + REFRESH"),
            ("Progress Router", "Activity Page",   "200 CompletionOut"),
            ("Activity Page",   "Redux Store",     "dispatch(fetchPairProgress(pairId)) — refresh dashboard"),
        ],
    },
    {
        "id": "SD-07",
        "title": "Speech Transcription — Whisper STT",
        "trigger": "User finishes recording and clicks Submit Recording",
        "route": "POST /api/speech/transcribe (multipart)",
        "participants": ["User", "AudioRecorder", "Browser MediaRecorder API", "Speech Router", "Whisper Service", "ffmpeg"],
        "steps": [
            ("User",              "AudioRecorder",         "click Record"),
            ("AudioRecorder",     "Browser MediaRecorder", "getUserMedia({audio:true}) → MediaStream"),
            ("Browser MediaRecorder", "AudioRecorder",     "dataavailable chunks every 250ms"),
            ("User",              "AudioRecorder",         "click Stop"),
            ("AudioRecorder",     "AudioRecorder",         "Blob(chunks, 'audio/webm')"),
            ("User",              "AudioRecorder",         "click Submit Recording"),
            ("AudioRecorder",     "Speech Router",         "POST /api/speech/transcribe (multipart FormData)"),
            ("Speech Router",     "Speech Router",         "validate MIME type → validate size (100B–25MB)"),
            ("Speech Router",     "Whisper Service",       "transcribe_audio(audio_bytes, filename)"),
            ("Whisper Service",   "Whisper Service",       "_load_model() [lazy, cached after first call]"),
            ("[ALT unavailable]", "",                      "Return {text:'', is_mock:True} → frontend shows Demo Mode warning"),
            ("Whisper Service",   "ffmpeg",                "transcode to 16kHz mono WAV"),
            ("Whisper Service",   "Whisper Service",       "_whisper_model.transcribe(out_wav, fp16=False)"),
            ("Whisper Service",   "Speech Router",         "{text, language, confidence, is_mock:False}"),
            ("Speech Router",     "AudioRecorder",         "200 TranscribeResponse"),
            ("AudioRecorder",     "User",                  "display transcription, enable Submit button"),
        ],
    },
    {
        "id": "SD-08",
        "title": "Friend Request Lifecycle",
        "trigger": "User A searches for and adds User B as a friend",
        "route": "POST /api/friends/request/{user_id}, PUT /api/friends/request/{req_id}/accept",
        "participants": ["User A", "User B", "React Frontend", "Friends Router", "PostgreSQL DB"],
        "steps": [
            ("User A",          "React Frontend",  "type username → GET /api/users/search?q={query}"),
            ("User A",          "React Frontend",  "click 'Add Friend' on User B's card"),
            ("React Frontend",  "Friends Router",  "POST /api/friends/request/{target_user_id}"),
            ("Friends Router",  "Friends Router",  "check sender != receiver (else 400)"),
            ("Friends Router",  "PostgreSQL DB",   "query existing FriendRequest between pair"),
            ("[ALT]", "",                           "Duplicate → 400 'Already sent or already friends'"),
            ("Friends Router",  "PostgreSQL DB",   "INSERT FriendRequest {pending}"),
            ("Friends Router",  "React Frontend",  "201 'Friend request sent'"),
            ("User B",          "React Frontend",  "GET /api/friends/requests → pending list"),
            ("[ALT accept]", "", "PUT /api/friends/request/{req_id}/accept → UPDATE status='accepted'"),
            ("[ALT decline]","", "PUT /api/friends/request/{req_id}/decline → UPDATE status='declined'"),
            ("Friends Router",  "React Frontend",  "200 Accepted / Declined"),
        ],
    },
    {
        "id": "SD-09",
        "title": "Leaderboard Load (Global + Friends)",
        "trigger": "User navigates to /leaderboard",
        "route": "GET /api/leaderboard/{pairId} and GET /api/leaderboard/{pairId}/friends",
        "participants": ["User", "LeaderboardPage", "Leaderboard Router", "PostgreSQL DB"],
        "steps": [
            ("User",               "LeaderboardPage",    "navigate('/leaderboard')"),
            ("[PAR]", "",           "Global and Friends fetches simultaneously"),
            ("LeaderboardPage",    "Leaderboard Router", "GET /api/leaderboard/{pairId}?limit=50"),
            ("Leaderboard Router", "PostgreSQL DB",      "JOIN UserLanguageProgress+User WHERE pair_id ORDER BY total_xp DESC LIMIT 50"),
            ("PostgreSQL DB",      "LeaderboardPage",    "List[LeaderboardEntry {rank, username, avatar, total_xp}]"),
            ("LeaderboardPage",    "Leaderboard Router", "GET /api/leaderboard/{pairId}/friends"),
            ("Leaderboard Router", "PostgreSQL DB",      "query accepted FriendRequests → collect friend_ids → JOIN ULP+User WHERE user_id IN friend_ids"),
            ("PostgreSQL DB",      "LeaderboardPage",    "List[LeaderboardEntry]"),
            ("LeaderboardPage",    "User",               "render Global + Friends tabs, highlight current user row"),
        ],
    },
    {
        "id": "SD-10",
        "title": "Admin: Create Language Pair",
        "trigger": "Admin fills new pair form and clicks Create",
        "route": "POST /api/admin/content/pairs",
        "participants": ["Admin", "AdminDashboard", "Admin Router", "Dependencies", "ContentService", "File System"],
        "steps": [
            ("Admin",          "AdminDashboard",  "fill pair_id form (e.g. 'hi-ko'), click Create"),
            ("AdminDashboard", "Admin Router",    "POST /api/admin/content/pairs {pair_id: 'hi-ko'}"),
            ("Admin Router",   "Dependencies",    "require_admin() — validate role == 'admin'"),
            ("Admin Router",   "ContentService",  "scaffold_pair('hi-ko')"),
            ("ContentService", "File System",     "LOOP: 3 months × 6 blocks × 8 types = 144 JSON files"),
            ("ContentService", "File System",     "CREATE meta.json"),
            ("ContentService", "Admin Router",    "success (144 files created)"),
            ("Admin Router",   "AdminDashboard",  "201 {message: 'Pair hi-ko created', files: 144}"),
            ("AdminDashboard", "Admin",           "refresh pairs list + show success notification"),
        ],
    },
    {
        "id": "SD-11",
        "title": "Server Startup Initialization",
        "trigger": "Uvicorn starts the FastAPI application",
        "route": "lifespan startup event",
        "participants": ["Uvicorn", "FastAPI App", "Database", "Security Module", "File System", "Logger"],
        "steps": [
            ("Uvicorn",     "FastAPI App", "startup event (lifespan)"),
            ("FastAPI App", "Database",    "create_tables() — import models → Base.metadata.create_all()"),
            ("Database",    "Database",    "CREATE tables IF NOT EXIST: users, user_language_progress, activity_completions, friend_requests"),
            ("FastAPI App", "FastAPI App", "seed_admin() — query admin user"),
            ("[ALT]", "",                  "Admin exists → LOG 'Admin already exists'"),
            ("[ALT]", "",                  "Admin missing → hash_password → INSERT admin user → COMMIT"),
            ("FastAPI App", "File System", "mkdir uploads/ if not exists"),
            ("FastAPI App", "FastAPI App", "ADD CORSMiddleware"),
            ("FastAPI App", "FastAPI App", "MOUNT /uploads → StaticFiles, /static → data/languages/"),
            ("FastAPI App", "FastAPI App", "REGISTER all routers under /api prefix"),
            ("FastAPI App", "Logger",      "LOG 'Startup complete. Server ready.'"),
            ("Uvicorn",     "Uvicorn",     "BIND 0.0.0.0:8000 — begin accepting HTTP connections"),
        ],
    },
]

for seq in sequences:
    add_heading(f"{seq['id']} — {seq['title']}", level=2)
    add_body(f"Trigger: {seq['trigger']}", italic=True)
    add_body(f"Route: {seq['route']}", italic=True)

    add_body("Participants:", bold=True)
    p = doc.add_paragraph()
    run = p.add_run("  ·  ".join(seq["participants"]))
    run.font.size = Pt(9)
    run.italic = True

    add_body("Message Flow:", bold=True)
    add_table(
        ["#", "From", "To", "Message / Action"],
        [
            (str(i + 1), frm, to, msg)
            for i, (frm, to, msg) in enumerate(seq["steps"])
        ],
        header_color="2E6EBF"
    )
    doc.add_paragraph()

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3: ACTIVITY DIAGRAMS
# ─────────────────────────────────────────────────────────────────────────────
add_heading("4. Section 3: Activity Diagrams", level=1)
add_note(
    "Activity diagrams use swimlanes to assign responsibility. [Decision] nodes are guards with "
    "YES/NO branches. fork/join notations indicate parallel execution. "
    "Each diagram covers a complete user journey end-to-end.",
    prefix="Notation: "
)
doc.add_paragraph()

activities = [
    {
        "id": "AD-01",
        "title": "User Registration Flow",
        "swimlanes": ["User", "Frontend", "Backend", "Database"],
        "steps": [
            ("User",             "START"),
            ("User",             "Fill registration form: {username, email, password}"),
            ("User",             "Click 'Register' button"),
            ("Frontend",         "Validate form locally: username 3–50 chars alphanumeric+underscore, password min 8, email format"),
            ("[Decision]",       "Form valid?"),
            ("Frontend",         "[NO] → Show inline validation errors → return to FILL form"),
            ("Frontend",         "[YES] → dispatch registerUser(credentials) to Redux"),
            ("Frontend/Redux",   "SET state.loading = true"),
            ("Frontend/Redux",   "CALL POST /api/auth/register via Axios"),
            ("Backend",          "RECEIVE POST /api/auth/register"),
            ("Backend",          "VALIDATE Pydantic schema (RegisterRequest)"),
            ("[Decision]",       "Schema valid?"),
            ("Backend",          "[NO] → RETURN 422 Unprocessable → Frontend SHOW error → END"),
            ("Backend",          "[YES] → QUERY users WHERE email == req.email"),
            ("[Decision]",       "Email already exists?"),
            ("Backend",          "[YES] → RETURN 400 'Email already registered' → END"),
            ("Backend",          "[NO] → QUERY users WHERE username == req.username"),
            ("[Decision]",       "Username already taken?"),
            ("Backend",          "[YES] → RETURN 400 'Username already taken' → END"),
            ("Backend",          "[NO] → HASH password (bcrypt)"),
            ("Backend",          "INSERT new User into database"),
            ("Backend",          "ISSUE JWT (HS256, 7-day expiry)"),
            ("Backend",          "RETURN 201 TokenResponse"),
            ("Database",         "COMMIT new User record"),
            ("Frontend/Redux",   "RECEIVE TokenResponse → SET state.token, user, isAuthenticated = true"),
            ("Frontend/Redux",   "WRITE lw_token + lw_user to localStorage"),
            ("Frontend/Redux",   "SET state.loading = false"),
            ("Frontend",         "NAVIGATE to /dashboard"),
            ("User",             "END"),
        ],
    },
    {
        "id": "AD-02",
        "title": "Dashboard Load and Render",
        "swimlanes": ["User", "DashboardPage", "Redux Store", "Backend", "File System"],
        "steps": [
            ("User",          "START → navigate('/dashboard')"),
            ("DashboardPage", "ProtectedRoute CHECK: isAuthenticated?"),
            ("[Decision]",    "Authenticated?"),
            ("DashboardPage", "[NO] → redirect to /login → END"),
            ("DashboardPage", "[YES] → render loading skeleton"),
            ("[FORK]",        "PARALLEL: dispatch(fetchAllProgress()) || dispatch(fetchPairs())"),
            ("Redux Store",   "GET /api/progress → RECEIVE List[ProgressOut] → SET allProgress"),
            ("Redux Store",   "GET /api/content/pairs → RECEIVE pair list → SET pairs"),
            ("[JOIN]",        "Both fetches complete"),
            ("DashboardPage", "AUTO-SELECT first pair if currentPairId is null"),
            ("[FORK]",        "PARALLEL: GET meta + GET completions"),
            ("Backend",       "GET /api/content/{pairId}/meta → read meta.json"),
            ("Backend",       "GET /api/progress/{pairId}/completions → List[CompletionOut]"),
            ("[JOIN]",        "Both fetches complete"),
            ("DashboardPage", "BUILD completedSeqIds = Set(completions WHERE passed=True)"),
            ("DashboardPage", "DETERMINE currentActivityId from progress (default 1)"),
            ("DashboardPage", "isUnlocked(seqId) = seqId <= currentActivityId"),
            ("DashboardPage", "isCompleted(seqId) = seqId IN completedSeqIds"),
            ("DashboardPage", "RENDER Month/Block/Activity roadmap"),
            ("[Decision]",    "For each activity node: isCompleted?"),
            ("DashboardPage", "[YES] → RENDER green ring + checkmark"),
            ("[Decision]",    "isUnlocked?"),
            ("DashboardPage", "[YES] → RENDER colored clickable node (by activity type)"),
            ("DashboardPage", "[NO]  → RENDER grey locked node with lock icon"),
            ("User",          "SEE roadmap with progress state → END"),
        ],
    },
    {
        "id": "AD-03",
        "title": "Activity Execution — General Flow",
        "swimlanes": ["User", "Activity Page", "Backend Validate", "Groq AI", "Scoring", "Backend Progress", "Database"],
        "steps": [
            ("User",             "START → click unlocked activity node"),
            ("Activity Page",    "ActivityRouter maps type → component"),
            ("Activity Page",    "GET /api/content/{pairId}/activity?file={activityFile}"),
            ("Activity Page",    "PARSE JSON → RENDER activity UI"),
            ("User",             "READ / LISTEN / SPEAK / WRITE answers"),
            ("[Decision]",       "All required answers filled?"),
            ("Activity Page",    "[NO] → SHOW 'Please answer all questions'"),
            ("Activity Page",    "[YES] → BUILD ValidateRequest"),
            ("Backend Validate", "POST /api/validate received"),
            ("[Decision]",       "activity_type in MCQ_TYPES ('test')?"),
            ("Scoring",          "[YES] → score_mcq_locally: str comparison per question → calculate_score"),
            ("Groq AI",          "[NO]  → _build_prompt(type, questions, max_xp, ...) → POST to LLaMA-3"),
            ("Groq AI",          "PARSE JSON response → merge user_answer/correct_answer"),
            ("Scoring",          "calculate_score: clamp scores → sum → percentage → passed = pct >= 20%"),
            ("Backend Validate", "_determine_feedback_tier(pct, attempt_count): hint/lesson/praise"),
            ("[Decision]",       "tier != 'lesson' AND GROQ type?"),
            ("Groq AI",          "[YES] → generate_tier_feedback → second Groq call"),
            ("[Decision]",       "SCORE_THRESHOLD_OVERRIDE > 0?"),
            ("Backend Validate", "[YES] → effective_passed = score >= OVERRIDE"),
            ("Backend Validate", "RETURN ValidateResponse"),
            ("Activity Page",    "SHOW ScoreModal {score, passed, XP, feedback, suggestion}"),
            ("Backend Progress", "POST /api/progress/{pairId}/complete"),
            ("[Decision]",       "First completion OR improvement?"),
            ("Backend Progress", "Award xp_delta accordingly"),
            ("[Decision]",       "passed=True AND seq_id == current_activity_id?"),
            ("Database",         "[YES] → INCREMENT current_activity_id → UNLOCK next activity"),
            ("Activity Page",    "dispatch(fetchPairProgress) → refresh dashboard"),
            ("User",             "SEE result popup → click Continue → back to Dashboard → END"),
        ],
    },
    {
        "id": "AD-04",
        "title": "STT Recording Flow (Speaking / Pronunciation)",
        "swimlanes": ["User", "AudioRecorder", "Browser API", "Backend Speech", "Whisper Service"],
        "steps": [
            ("User",           "START → click 'Start Recording'"),
            ("AudioRecorder",  "REQUEST microphone permission"),
            ("[Decision]",     "Permission granted?"),
            ("AudioRecorder",  "[NO] → SHOW 'Microphone access denied' → END"),
            ("AudioRecorder",  "[YES] → INITIALIZE MediaRecorder(stream, 'audio/webm')"),
            ("AudioRecorder",  "mediaRecorder.start() → START red pulse indicator"),
            ("Browser API",    "EMIT dataavailable events every 250ms"),
            ("AudioRecorder",  "COLLECT audio chunks array"),
            ("User",           "SPEAKS into microphone"),
            ("User",           "CLICK 'Stop Recording'"),
            ("AudioRecorder",  "mediaRecorder.stop() → final chunk received"),
            ("AudioRecorder",  "COMBINE chunks: new Blob(chunks, 'audio/webm')"),
            ("AudioRecorder",  "CREATE object URL → render <audio> preview controls"),
            ("User",           "LISTENS to preview → click 'Submit Recording'"),
            ("AudioRecorder",  "CREATE FormData, append Blob as file"),
            ("Backend Speech", "POST /api/speech/transcribe (multipart FormData)"),
            ("Backend Speech", "VALIDATE MIME type (startswith checks)"),
            ("Backend Speech", "READ audio bytes → VALIDATE size 100B < size < 25MB"),
            ("[Decision]",     "Whisper model available?"),
            ("Backend Speech", "[NO] → RETURN {is_mock:True} → SHOW 'Demo mode' warning → block submit → END"),
            ("Whisper Service","[YES] → WRITE bytes to tempfile"),
            ("Whisper Service","TRANSCODE via ffmpeg to 16kHz mono WAV"),
            ("Whisper Service","whisper_model.transcribe(wav_path, fp16=False)"),
            ("Whisper Service","EXTRACT text, language, confidence → CLEANUP temp files"),
            ("Whisper Service","RETURN {text, language, confidence, is_mock:False}"),
            ("AudioRecorder",  "SET user_answer = transcribed text"),
            ("User",           "REVIEW transcription → CLICK 'Submit Answers' → standard Validate flow → END"),
        ],
    },
    {
        "id": "AD-05",
        "title": "Friend Request System Flow",
        "swimlanes": ["User A (Sender)", "User B (Receiver)", "Frontend", "Backend Friends Router", "Database"],
        "steps": [
            ("User A",               "START → navigate to /search"),
            ("User A",               "TYPE username in search box"),
            ("Frontend",             "GET /api/users/search?q={query} → display matching user cards"),
            ("User A",               "CLICK 'Add Friend' on User B's card"),
            ("Backend Friends Router","[Decision] Is self-request?"),
            ("Backend Friends Router","[YES] → RETURN 400 → END"),
            ("Backend Friends Router","[NO] → QUERY User B exists AND is_active"),
            ("[Decision]",           "User B found?"),
            ("Backend Friends Router","[NO] → RETURN 404 → END"),
            ("Backend Friends Router","QUERY existing FriendRequest (either direction)"),
            ("[Decision]",           "Request already exists?"),
            ("Backend Friends Router","[YES] → RETURN 400 'Already sent/friends' → END"),
            ("Database",             "[NO] → INSERT FriendRequest {sender=A, receiver=B, status='pending'}"),
            ("Frontend",             "SHOW 'Request sent' toast to User A"),
            ("User B",               "NAVIGATE to /friends or /search"),
            ("Frontend",             "GET /api/friends/requests → display pending requests"),
            ("[Decision]",           "User B: Accept or Decline?"),
            ("Backend Friends Router","[ACCEPT] → PUT /api/friends/request/{id}/accept → UPDATE status='accepted'"),
            ("Frontend",             "SHOW 'Now friends!' confirmation"),
            ("Backend Friends Router","[DECLINE] → PUT /api/friends/request/{id}/decline → UPDATE status='declined'"),
            ("Frontend",             "REMOVE request from list → END"),
        ],
    },
    {
        "id": "AD-06",
        "title": "Progress Advancement & XP Unlock Logic",
        "swimlanes": ["Progress Router", "Scoring Logic", "Database"],
        "steps": [
            ("Progress Router", "START → RECEIVE CompleteActivityRequest"),
            ("[Decision]",      "SCORE_THRESHOLD_OVERRIDE > 0?"),
            ("Progress Router", "[YES] → effective_passed = score_earned >= OVERRIDE"),
            ("Progress Router", "[NO]  → effective_passed = req.passed"),
            ("Progress Router", "DERIVE month/block from seq_id: zero_based=id-1, month=(zero//48)+1, block=((zero%48)//8)+1"),
            ("Database",        "QUERY UserLanguageProgress WHERE user_id AND lang_pair_id"),
            ("[Decision]",      "Progress record exists?"),
            ("Database",        "[NO] → INSERT new UserLanguageProgress (m=1, b=1, a=seq_id)"),
            ("Database",        "QUERY ActivityCompletion WHERE user_id AND lang_pair_id AND seq_id"),
            ("[Decision]",      "Completion already exists (retry)?"),
            ("[Decision]",      "[YES retry] → new_score > existing.score_earned?"),
            ("Progress Router", "  [YES] → xp_delta = new - old → UPDATE score_earned, passed"),
            ("Progress Router", "  [NO]  → xp_delta = 0 (no XP for same/lower score)"),
            ("Database",        "  UPDATE attempts += 1, ai_feedback, completed_at"),
            ("Database",        "[NO first] → INSERT new ActivityCompletion, xp_delta = score_earned"),
            ("Database",        "UPDATE total_xp += xp_delta"),
            ("Database",        "UPDATE last_activity_at = utcnow()"),
            ("[Decision]",      "effective_passed == True AND seq_id == current_activity_id?"),
            ("Database",        "[YES] → next_id = current_activity_id + 1"),
            ("Database",        "        UPDATE current_activity_id = next_id"),
            ("Database",        "        UPDATE current_month, current_block = _derive(next_id)"),
            ("Progress Router", "NOTE: Next activity is now UNLOCKED on the dashboard"),
            ("Database",        "COMMIT all changes → REFRESH completion → RETURN CompletionOut → END"),
        ],
    },
    {
        "id": "AD-07",
        "title": "Admin Curriculum Management Flow",
        "swimlanes": ["Admin", "AdminDashboard", "Backend Admin Router", "ContentService", "File System"],
        "steps": [
            ("Admin",                 "START → login via /admin/login"),
            ("AdminDashboard",        "GET /api/admin/content/pairs → display pairs"),
            ("AdminDashboard",        "GET /api/admin/analytics → display stats"),
            ("Admin",                 "[ACTION A] CREATE NEW LANGUAGE PAIR"),
            ("Admin",                 "ENTER pair_id (e.g. 'hi-ja'), click Create"),
            ("Backend Admin Router",  "require_admin() → CALL ContentService.scaffold_pair('hi-ja')"),
            ("ContentService",        "LOOP Month 1-3 × Block 1-6 × 8 types → CREATE M{m}B{b}_{type}.json"),
            ("File System",           "WRITE 144 JSON files + meta.json"),
            ("AdminDashboard",        "REFRESH pairs list → show new pair"),
            ("Admin",                 "[ACTION B] EDIT ACTIVITY CONTENT"),
            ("Admin",                 "SELECT pair → month → block → activity type"),
            ("ContentService",        "GET /api/admin/content/{pair}/{m}/{b}/{type} → return JSON"),
            ("AdminDashboard",        "Render JSON editor/form → Admin edits content"),
            ("Admin",                 "CLICK Save"),
            ("ContentService",        "PUT → OVERWRITE activity .json file on disk"),
            ("Admin",                 "[ACTION C] MANAGE USERS"),
            ("[Decision]",            "Activate / Deactivate / Promote?"),
            ("[Decision]",            "Trying to deactivate own admin account?"),
            ("Backend Admin Router",  "[YES] → RETURN 400 'Cannot deactivate yourself' → END"),
            ("Database",              "[NO]  → UPDATE user.is_active or role"),
            ("Admin",                 "END"),
        ],
    },
    {
        "id": "AD-08",
        "title": "Onboarding Flow (New User First Login)",
        "swimlanes": ["User", "Frontend", "Backend Progress", "Database"],
        "steps": [
            ("User",            "START → complete registration"),
            ("Frontend",        "POST-REGISTER redirect to /onboarding"),
            ("Frontend",        "GET /api/content/pairs → fetch available language pairs"),
            ("Frontend",        "DISPLAY pair selection cards (e.g. Hindi → Japanese)"),
            ("User",            "SELECT a language pair, click 'Start Learning'"),
            ("Frontend",        "dispatch startLanguagePair(pairId)"),
            ("Backend Progress","POST /api/progress/{pairId}/start"),
            ("Database",        "QUERY UserLanguageProgress WHERE user_id AND pair_id"),
            ("[Decision]",      "Already started?"),
            ("Backend Progress","[YES] → RETURN existing ProgressOut (idempotent)"),
            ("Database",        "[NO]  → INSERT UserLanguageProgress {xp=0, month=1, block=1, activity_id=1} → COMMIT"),
            ("Frontend",        "SET currentPairId in Redux"),
            ("Frontend",        "WRITE 'lw_pair' to localStorage"),
            ("Frontend",        "NAVIGATE to /dashboard"),
            ("User",            "SEE roadmap: Activity #1 (Lesson) UNLOCKED, all others LOCKED → END"),
        ],
    },
]

for act in activities:
    add_heading(f"{act['id']} — {act['title']}", level=2)
    add_body(f"Swimlanes: {', '.join(act['swimlanes'])}", italic=True)
    doc.add_paragraph()
    add_table(
        ["Step", "Swimlane / Actor", "Action / Decision Node"],
        [
            (str(i + 1), lane, action)
            for i, (lane, action) in enumerate(act["steps"])
        ],
        header_color="1A6B3E"
    )
    doc.add_paragraph()

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4: STATE DIAGRAMS
# ─────────────────────────────────────────────────────────────────────────────
add_heading("5. Section 4: State Diagrams", level=1)
add_note(
    "State diagrams document all stateful entities in the system. Each machine lists: "
    "all discrete states, all transitions (event [guard] / action → TARGET_STATE), "
    "entry/exit actions, and composite sub-states where applicable.",
    prefix="Notation: "
)
doc.add_paragraph()

state_machines = [
    {
        "id": "SM-01",
        "title": "User Account Lifecycle",
        "entity": "User (database record)",
        "files": "backend/app/models/user.py · backend/app/routers/auth.py · backend/app/routers/admin.py",
        "states": [
            ("UNREGISTERED", "Initial state — no account exists"),
            ("REGISTERED_ACTIVE", "entry: created_at=utcnow(), is_active=True, role='user'"),
            ("DEACTIVATED", "entry: is_active=False  |  exit: is_active=True"),
            ("LOGGED_IN", "entry: last_active=utcnow(), JWT issued"),
            ("LOGGED_OUT", "entry: localStorage.removeItem('lw_token'), removeItem('lw_user')"),
            ("ADMIN_ROLE", "Sub-state of REGISTERED_ACTIVE for role='admin' users"),
        ],
        "transitions": [
            ("[*] → UNREGISTERED", "", "Initial pseudo-state"),
            ("UNREGISTERED → REGISTERED_ACTIVE", "POST /api/auth/register [valid form, unique credentials]", ""),
            ("REGISTERED_ACTIVE → DEACTIVATED",  "admin calls deactivate", "is_active = False"),
            ("REGISTERED_ACTIVE → LOGGED_IN",    "POST /api/auth/login [valid password]", ""),
            ("REGISTERED_ACTIVE → [*]",          "DELETE user cascade", "account deleted"),
            ("DEACTIVATED → REGISTERED_ACTIVE",  "admin calls activate", "is_active = True"),
            ("LOGGED_IN → LOGGED_OUT",           "client calls logout OR JWT expires/invalid", ""),
            ("LOGGED_IN → DEACTIVATED",          "admin deactivates account", ""),
            ("LOGGED_OUT → LOGGED_IN",           "POST /api/auth/login [valid credentials]", ""),
            ("LOGGED_IN → LOGGED_IN",            "GET /api/auth/me [JWT valid]", "self-loop"),
        ],
    },
    {
        "id": "SM-02",
        "title": "JWT Token Lifecycle",
        "entity": "JWT Bearer Token",
        "files": "backend/app/core/security.py · src/api/client.js",
        "states": [
            ("DOES_NOT_EXIST", "No token stored"),
            ("ACTIVE", "entry: exp=utcnow()+7days, stored in localStorage as 'lw_token'"),
            ("VALID", "Token decoded successfully — payload returned to dependencies"),
            ("INVALID", "JWTError raised during decode (bad signature/format)"),
            ("EXPIRED", "Current time > exp claim — decode fails"),
            ("CLEARED", "entry: localStorage.removeItem('lw_token'), removeItem('lw_user')"),
            ("ACCEPTED", "Request proceeds to endpoint handler"),
            ("REJECTED", "Returns HTTPException 401 to client"),
        ],
        "transitions": [
            ("[*] → DOES_NOT_EXIST", "Initial", ""),
            ("DOES_NOT_EXIST → ACTIVE", "successful login/register", "token stored"),
            ("ACTIVE → VALID",          "decode_token [signature ok, not expired]", ""),
            ("ACTIVE → INVALID",        "JWTError raised",  ""),
            ("ACTIVE → EXPIRED",        "time > exp",        ""),
            ("VALID → ACCEPTED",        "get_current_user [user found, is_active=True]", ""),
            ("VALID → REJECTED",        "get_current_user [user not found or inactive]", ""),
            ("INVALID → CLEARED",       "Axios receives 401", ""),
            ("EXPIRED → CLEARED",       "Axios receives 401", ""),
            ("CLEARED → DOES_NOT_EXIST","", "redirect to /login or /admin/login"),
        ],
    },
    {
        "id": "SM-03",
        "title": "UserLanguageProgress State Machine",
        "entity": "UserLanguageProgress (database record)",
        "files": "backend/app/models/progress.py · backend/app/routers/progress.py",
        "states": [
            ("NOT_STARTED", "No progress record for this pair yet"),
            ("ACTIVE", "entry: started_at=utcnow()  — composite state containing sub-states"),
            ("AT_ACTIVITY(n)", "sub-state: current position in curriculum, n=1 to 144"),
            ("BEGINNER", "sub-state: total_xp < 500"),
            ("INTERMEDIATE", "sub-state: total_xp 500–2000"),
            ("ADVANCED", "sub-state: total_xp > 2000"),
            ("CURRICULUM_COMPLETE", "All 144 activities passed — user can still retry for XP"),
        ],
        "transitions": [
            ("[*] → NOT_STARTED", "Initial", ""),
            ("NOT_STARTED → ACTIVE", "POST /api/progress/{pair_id}/start", "entry: xp=0, month=1, block=1, activity_id=1"),
            ("AT_ACTIVITY(n) → AT_ACTIVITY(n) [self-loop]", "POST /complete [passed=False OR seq_id!=current]", "XP delta only, no position advance"),
            ("AT_ACTIVITY(n) → AT_ACTIVITY(n+1)", "POST /complete [passed=True AND seq_id==current]", "current_activity_id++, re-derive month/block"),
            ("AT_ACTIVITY(144) → CURRICULUM_COMPLETE", "POST /complete [passed=True]", ""),
            ("CURRICULUM_COMPLETE → ACTIVE [self-loop]", "retry any past activity", "XP improvement only, no position advance"),
        ],
    },
    {
        "id": "SM-04",
        "title": "ActivityCompletion Record State",
        "entity": "ActivityCompletion (database record)",
        "files": "backend/app/models/progress.py · backend/app/routers/progress.py",
        "states": [
            ("NEVER_ATTEMPTED", "No completion record exists for this activity"),
            ("ATTEMPTED", "Record exists; user can retry"),
            ("FAILED", "sub-state: passed=False"),
            ("PASSED", "sub-state: passed=True  — XP improvement possible on retry"),
        ],
        "transitions": [
            ("[*] → NEVER_ATTEMPTED", "Initial", ""),
            ("NEVER_ATTEMPTED → ATTEMPTED", "POST /complete [first attempt]", "INSERT record, attempts=1"),
            ("FAILED → PASSED", "POST /complete [passes, new_score > old_score]", "xp_delta = new-old, UPDATE score"),
            ("PASSED → PASSED [self-loop]", "POST /complete [retry, new_score > old]", "xp_delta = new-old, attempts++"),
            ("PASSED → PASSED [self-loop]", "POST /complete [retry, new_score <= old]", "no XP, attempts++"),
        ],
    },
    {
        "id": "SM-05",
        "title": "FriendRequest State Machine",
        "entity": "FriendRequest (database record)",
        "files": "backend/app/models/friends.py · backend/app/routers/friends.py",
        "states": [
            ("DOES_NOT_EXIST", "No friend relation between User A and User B"),
            ("PENDING", "entry: status='pending', created_at=utcnow()"),
            ("ACCEPTED", "entry: status='accepted' — both users appear on each other's friends leaderboard"),
            ("DECLINED", "entry: status='declined' — terminal state in current implementation"),
            ("REMOVED", "Friendship deleted — row removed from DB"),
        ],
        "transitions": [
            ("[*] → DOES_NOT_EXIST", "Initial", ""),
            ("DOES_NOT_EXIST → PENDING", "POST /friends/request/{user_id} [target exists, not self, no duplicate]", "INSERT FriendRequest"),
            ("PENDING → ACCEPTED", "PUT /friends/request/{id}/accept [caller==receiver_id]", "UPDATE status='accepted'"),
            ("PENDING → DECLINED", "PUT /friends/request/{id}/decline [caller==receiver_id]", "UPDATE status='declined'"),
            ("ACCEPTED → REMOVED", "DELETE /friends/{user_id} [caller is sender OR receiver]", "DELETE row from DB"),
            ("DECLINED → [*]", "Terminal state", ""),
            ("REMOVED → DOES_NOT_EXIST", "row deleted — new request can be sent", ""),
        ],
    },
    {
        "id": "SM-06",
        "title": "Redux Auth State (Frontend)",
        "entity": "Redux authSlice state",
        "files": "src/store/authSlice.js",
        "states": [
            ("UNAUTHENTICATED", "entry: token=null, user=null, isAuthenticated=false"),
            ("LOADING", "entry: loading=true, error=null"),
            ("AUTHENTICATED", "entry: token=JWT, user=UserOut, isAuthenticated=true, localStorage written"),
            ("ERROR", "entry: loading=false, error='error message string'"),
        ],
        "transitions": [
            ("[*] → UNAUTHENTICATED", "Initial OR localStorage has lw_token → hydrate to AUTHENTICATED", ""),
            ("UNAUTHENTICATED → LOADING", "dispatch loginUser.pending OR registerUser.pending", ""),
            ("LOADING → AUTHENTICATED", "loginUser.fulfilled OR registerUser.fulfilled", ""),
            ("LOADING → ERROR", "loginUser.rejected OR registerUser.rejected", ""),
            ("AUTHENTICATED → UNAUTHENTICATED", "dispatch logout OR Axios 401 interceptor fires", ""),
            ("AUTHENTICATED → AUTHENTICATED [self-loop]", "dispatch updateUser(patch)", "user merged with patch"),
            ("ERROR → LOADING", "user retries login", ""),
            ("ERROR → UNAUTHENTICATED", "dispatch clearError", ""),
        ],
    },
    {
        "id": "SM-07",
        "title": "Activity Page State (Generic)",
        "entity": "Individual Activity Page React component",
        "files": "src/pages/activities/ (LessonPage, WritingPage, SpeakingPage, etc.)",
        "states": [
            ("LOADING_CONTENT", "entry: fetch GET /api/content/{pairId}/activity"),
            ("LOAD_ERROR", "Show error message with ← Back button"),
            ("CONTENT_READY", "Activity rendered with instructions + questions"),
            ("IDLE", "sub-state: user reading content, no answers yet"),
            ("ANSWERING", "sub-state: one or more answers entered"),
            ("SUBMITTING", "entry: POST /api/validate payload, loading spinner"),
            ("PASSED_RESULT", "entry: ScoreModal with green pass indicator, XP earned"),
            ("FAILED_RESULT", "entry: ScoreModal with red fail indicator, suggestion shown"),
            ("RECORDING_ACTIVE", "nested sub-state (Speaking/Pronunciation): MediaRecorder.start()"),
            ("PROCESSING_AUDIO", "nested sub-state: POST /api/speech/transcribe in progress"),
            ("TRANSCRIBED", "nested sub-state: user_answer = transcribed text"),
            ("MOCK_WARNING", "nested sub-state: Whisper unavailable, submit blocked"),
        ],
        "transitions": [
            ("[*] → LOADING_CONTENT", "component mounted", ""),
            ("LOADING_CONTENT → CONTENT_READY", "fetch success", "activityData parsed"),
            ("LOADING_CONTENT → LOAD_ERROR", "fetch error", ""),
            ("LOAD_ERROR → [*]", "user clicks ← Back", "navigate to /dashboard"),
            ("IDLE → ANSWERING", "user starts filling answers", ""),
            ("ANSWERING → SUBMITTING", "all answers complete + user clicks Submit", ""),
            ("SUBMITTING → PASSED_RESULT", "validate success, passed=True", ""),
            ("SUBMITTING → FAILED_RESULT", "validate success, passed=False", ""),
            ("SUBMITTING → ANSWERING", "network error", "show error toast"),
            ("PASSED_RESULT → [*]", "user clicks Continue", "navigate to /dashboard"),
            ("PASSED_RESULT → CONTENT_READY", "user clicks Retry", "reset state"),
            ("FAILED_RESULT → CONTENT_READY", "user clicks Try Again", "attempt_count++"),
            ("NOT_RECORDING → RECORDING_ACTIVE", "user clicks Record", "MediaRecorder.start()"),
            ("RECORDING_ACTIVE → PROCESSING_AUDIO", "user clicks Stop", ""),
            ("PROCESSING_AUDIO → TRANSCRIBED", "is_mock=False, text returned", "user_answer = text"),
            ("PROCESSING_AUDIO → MOCK_WARNING", "is_mock=True", "submit blocked"),
        ],
    },
    {
        "id": "SM-08",
        "title": "Whisper Service Availability State",
        "entity": "WhisperService Python module",
        "files": "backend/app/services/whisper_service.py",
        "states": [
            ("UNINITIALIZED", "_load_model() not yet called; _whisper_available=None, _whisper_model=None"),
            ("LOADING", "entry: _load_model() called; import whisper; whisper.load_model()"),
            ("AVAILABLE", "entry: _whisper_model cached; _whisper_available=True"),
            ("UNAVAILABLE", "entry: _whisper_available=False — all calls return is_mock=True without retry"),
            ("TRANSCRIBING", "sub-state of AVAILABLE: ffmpeg transcode → whisper_model.transcribe()"),
        ],
        "transitions": [
            ("[*] → UNINITIALIZED", "Initial", ""),
            ("UNINITIALIZED → LOADING", "first call to transcribe_audio()", ""),
            ("LOADING → AVAILABLE", "whisper.load_model() success", "_whisper_model cached"),
            ("LOADING → UNAVAILABLE", "ImportError or any exception during load", "_whisper_available=False"),
            ("AVAILABLE → TRANSCRIBING", "transcribe_audio() called", ""),
            ("TRANSCRIBING → AVAILABLE", "success", "return {text, language, confidence, is_mock:False}"),
            ("TRANSCRIBING → AVAILABLE", "exception during transcription", "return {text:'', is_mock:False}"),
            ("UNAVAILABLE → UNAVAILABLE [self-loop]", "transcribe_audio() called", "return {text:'', is_mock:True}"),
        ],
    },
    {
        "id": "SM-09",
        "title": "Groq AI Service Call State",
        "entity": "Single Groq API call in GroqService",
        "files": "backend/app/services/groq_service.py",
        "states": [
            ("IDLE", "Waiting for validate_activity() or generate_tier_feedback() call"),
            ("BUILDING_PROMPT", "Composing system rubric + romanization rules + per-question input"),
            ("CALLING_GROQ", "entry: client.chat.completions.create(model=llama3-8b-8192, temperature=0.1)"),
            ("PARSING_RESPONSE", "Extract JSON, merge user_answer/correct_answer from original questions"),
            ("RESULT_READY", "Parsed result dict ready to return to caller"),
            ("ERROR", "Timeout, network error, or malformed JSON — fallback dict returned"),
        ],
        "transitions": [
            ("[*] → IDLE", "Initial", ""),
            ("IDLE → BUILDING_PROMPT", "validate_activity() or generate_tier_feedback() called", ""),
            ("BUILDING_PROMPT → CALLING_GROQ", "prompt complete", ""),
            ("CALLING_GROQ → PARSING_RESPONSE", "API responds with valid JSON", ""),
            ("CALLING_GROQ → ERROR", "API timeout or network error", ""),
            ("PARSING_RESPONSE → RESULT_READY", "JSON parse success", ""),
            ("PARSING_RESPONSE → ERROR", "malformed JSON", "return fallback {overall_feedback:'AI unavailable'}"),
            ("RESULT_READY → IDLE", "return dict to caller", ""),
            ("ERROR → IDLE", "return fallback dict", "generate_tier_feedback errors silently fall back"),
        ],
    },
    {
        "id": "SM-10",
        "title": "AudioRecorder Component State",
        "entity": "AudioRecorder React component",
        "files": "src/components/AudioRecorder.jsx",
        "states": [
            ("READY", "entry: audioBlob=null, chunks=[], is_recording=false"),
            ("REQUESTING_PERMISSION", "entry: navigator.mediaDevices.getUserMedia({audio:true})"),
            ("PERMISSION_DENIED", "entry: show 'Microphone access denied' error"),
            ("RECORDING", "entry: MediaRecorder.start(), red pulse animation shown"),
            ("PROCESSING_LOCAL", "entry: mediaRecorder.stop()"),
            ("PREVIEW", "entry: <audio src={objectURL} controls> rendered for playback"),
            ("UPLOADING", "entry: FormData append, POST /api/speech/transcribe, spinner shown"),
            ("TRANSCRIBED", "parent ActivityPage receives text as user_answer"),
            ("MOCK_FALLBACK", "entry: 'Whisper unavailable — demo mode' warning, Submit blocked"),
            ("UPLOAD_ERROR", "entry: 'Upload failed, please try again' toast"),
        ],
        "transitions": [
            ("[*] → READY", "Initial", ""),
            ("READY → REQUESTING_PERMISSION", "user clicks Record", ""),
            ("REQUESTING_PERMISSION → RECORDING", "permission granted", ""),
            ("REQUESTING_PERMISSION → PERMISSION_DENIED", "permission denied", ""),
            ("PERMISSION_DENIED → READY", "user acknowledges error", ""),
            ("RECORDING → RECORDING [self-loop]", "dataavailable chunk event", "push chunk to chunks[]"),
            ("RECORDING → PROCESSING_LOCAL", "user clicks Stop", ""),
            ("PROCESSING_LOCAL → PREVIEW", "stop event fired", "blob created, objectURL generated"),
            ("PREVIEW → READY", "user clicks Re-record", "URL.revokeObjectURL, reset chunks"),
            ("PREVIEW → UPLOADING", "user clicks Submit Recording", ""),
            ("UPLOADING → TRANSCRIBED", "response is_mock=False, text returned", "pass text to parent"),
            ("UPLOADING → MOCK_FALLBACK", "response is_mock=True", ""),
            ("UPLOADING → UPLOAD_ERROR", "network error", ""),
            ("TRANSCRIBED → READY", "reset for re-recording", ""),
            ("MOCK_FALLBACK → READY", "user clicks Re-record", ""),
            ("UPLOAD_ERROR → UPLOADING", "user clicks Retry", ""),
            ("UPLOAD_ERROR → READY", "user clicks Re-record", ""),
        ],
    },
    {
        "id": "SM-11",
        "title": "Frontend Route Guard Lifecycle",
        "entity": "ProtectedRoute and AdminRoute components",
        "files": "src/components/ProtectedRoute.jsx · src/components/AdminRoute.jsx",
        "states": [
            ("CHECKING_AUTH", "ProtectedRoute: read Redux state.auth.isAuthenticated"),
            ("AUTHORIZED", "ProtectedRoute: render children (MainLayout + Outlet)"),
            ("UNAUTHORIZED", "ProtectedRoute: Navigate to /login (replace:true)"),
            ("CHECKING_ADMIN", "AdminRoute: read isAuthenticated AND user.role"),
            ("ADMIN_AUTHORIZED", "AdminRoute: render <Outlet /> (AdminDashboard)"),
            ("ADMIN_UNAUTHORIZED", "AdminRoute: Navigate to /admin/login (replace:true)"),
        ],
        "transitions": [
            ("[*] → CHECKING_AUTH / CHECKING_ADMIN", "route change", ""),
            ("CHECKING_AUTH → AUTHORIZED",       "isAuthenticated == true", ""),
            ("CHECKING_AUTH → UNAUTHORIZED",     "isAuthenticated == false", ""),
            ("AUTHORIZED → UNAUTHORIZED",        "JWT expires + 401 OR dispatch logout", ""),
            ("UNAUTHORIZED → [*]",               "redirect to /login, no back history entry", ""),
            ("CHECKING_ADMIN → ADMIN_AUTHORIZED","isAuthenticated AND role=='admin'", ""),
            ("CHECKING_ADMIN → ADMIN_UNAUTHORIZED","not admin or not authenticated", ""),
            ("ADMIN_AUTHORIZED → ADMIN_UNAUTHORIZED", "JWT expires + 401", ""),
            ("ADMIN_UNAUTHORIZED → [*]",         "redirect to /admin/login", ""),
        ],
    },
    {
        "id": "SM-12",
        "title": "Redux Progress State",
        "entity": "Redux progressSlice state",
        "files": "src/store/progressSlice.js",
        "states": [
            ("EMPTY", "entry: allProgress=[], pairs=[], currentPairId=null, loading=false"),
            ("LOADING_PROGRESS", "entry: loading=true — fetchAllProgress pending"),
            ("LOADING_PAIRS", "entry: fetchPairs pending"),
            ("PAIRS_LOADED", "pairs array populated"),
            ("LOADED", "entry: loading=false — all progress data available"),
        ],
        "transitions": [
            ("[*] → EMPTY", "Initial", ""),
            ("EMPTY → LOADING_PROGRESS", "dispatch fetchAllProgress.pending", ""),
            ("EMPTY → LOADING_PAIRS", "dispatch fetchPairs.pending", ""),
            ("LOADING_PROGRESS → LOADED", "fetchAllProgress.fulfilled", "allProgress = payload"),
            ("LOADING_PROGRESS → EMPTY", "fetchAllProgress.rejected", "loading=false"),
            ("LOADING_PAIRS → PAIRS_LOADED", "fetchPairs.fulfilled", "pairs = payload"),
            ("PAIRS_LOADED → LOADED", "all data ready", ""),
            ("LOADED → LOADED [self-loop]", "dispatch setCurrentPair / updateProgressXP / advanceProgress / fetchPairProgress.fulfilled", "state updated"),
            ("LOADED → LOADING_PROGRESS", "dispatch fetchAllProgress again", "refresh cycle"),
        ],
    },
    {
        "id": "SM-13",
        "title": "Admin Dashboard Page State",
        "entity": "AdminDashboard React component",
        "files": "src/pages/admin/AdminDashboard.jsx",
        "states": [
            ("LOADING_ADMIN_DATA", "entry: GET /api/admin/users, GET /api/admin/content/pairs, GET /api/admin/analytics"),
            ("VIEWING_DASHBOARD", "entry: render user list, pair list, stats panels"),
            ("ACCESS_DENIED", "any fetch returns 403 — redirect to /admin/login"),
            ("USERS_TAB", "viewing user management section"),
            ("MODIFYING_USER", "entry: POST activate or deactivate API call in progress"),
            ("CURRICULUM_TAB", "viewing content management section"),
            ("PAIR_SELECTED", "a language pair is selected"),
            ("EDITING_ACTIVITY", "entry: GET activity JSON, render editor/form"),
            ("SAVING_ACTIVITY", "entry: PUT API call to write JSON to disk"),
            ("CREATING_PAIR", "entry: POST /api/admin/content/pairs"),
            ("DELETING_PAIR", "entry: DELETE /api/admin/content/pairs/{id}"),
        ],
        "transitions": [
            ("[*] → LOADING_ADMIN_DATA", "page mounted", ""),
            ("LOADING_ADMIN_DATA → VIEWING_DASHBOARD", "all fetches complete", ""),
            ("LOADING_ADMIN_DATA → ACCESS_DENIED", "fetch returns 403", ""),
            ("ACCESS_DENIED → [*]", "redirect to /admin/login", ""),
            ("VIEWING_DASHBOARD → USERS_TAB", "admin switches to users tab", ""),
            ("USERS_TAB → MODIFYING_USER", "admin clicks Activate/Deactivate", ""),
            ("MODIFYING_USER → USERS_TAB", "API success or error", "refresh user list"),
            ("VIEWING_DASHBOARD → CURRICULUM_TAB", "admin switches to curriculum tab", ""),
            ("CURRICULUM_TAB → PAIR_SELECTED", "admin clicks a pair", ""),
            ("PAIR_SELECTED → EDITING_ACTIVITY", "admin drills to activity type", ""),
            ("EDITING_ACTIVITY → SAVING_ACTIVITY", "admin clicks Save", ""),
            ("SAVING_ACTIVITY → PAIR_SELECTED", "success → 'Saved' toast", ""),
            ("SAVING_ACTIVITY → EDITING_ACTIVITY", "error → show error toast", ""),
            ("CURRICULUM_TAB → CREATING_PAIR", "click Create New Pair", ""),
            ("CREATING_PAIR → CURRICULUM_TAB", "success or error", "refresh pairs"),
            ("CURRICULUM_TAB → DELETING_PAIR", "click Delete Pair", ""),
            ("DELETING_PAIR → CURRICULUM_TAB", "success or error", "refresh pairs"),
        ],
    },
]

for sm in state_machines:
    add_heading(f"{sm['id']} — {sm['title']}", level=2)
    add_body(f"Entity: {sm['entity']}", bold=True)
    add_body(f"Files: {sm['files']}", italic=True)
    doc.add_paragraph()

    add_body("States:", bold=True)
    add_table(
        ["State Name", "Description / Entry Action"],
        sm["states"],
        header_color="7B2D8B"
    )

    add_body("Transitions:", bold=True)
    add_table(
        ["Transition", "Event [Guard]", "Action"],
        sm["transitions"],
        header_color="7B2D8B"
    )
    doc.add_paragraph()

page_break()


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5: ENTITY REFERENCE TABLE
# ─────────────────────────────────────────────────────────────────────────────
add_heading("6. Complete Stateful Entity Reference Table", level=1)
add_body(
    "The following table provides a complete cross-reference of all stateful entities in the "
    "LearnWise platform, mapping each entity to its State Machine diagram section, "
    "key states, and triggering events."
)
doc.add_paragraph()

add_table(
    ["Entity", "SM #", "Key States", "Trigger Events"],
    [
        ["User (DB model)",              "SM-01", "UNREGISTERED → REGISTERED_ACTIVE → DEACTIVATED ↔ LOGGED_IN", "register, login, logout, admin activate/deactivate"],
        ["JWT Token",                    "SM-02", "DOES_NOT_EXIST → ACTIVE → EXPIRED/INVALID → CLEARED",         "login, token decode, 401 response"],
        ["UserLanguageProgress (DB)",    "SM-03", "NOT_STARTED → ACTIVE (AT_ACTIVITY_N) → CURRICULUM_COMPLETE",  "start pair, complete activity (passed)"],
        ["ActivityCompletion (DB)",      "SM-04", "NEVER_ATTEMPTED → ATTEMPTED (FAILED / PASSED)",                "POST /progress/complete"],
        ["FriendRequest (DB)",           "SM-05", "DOES_NOT_EXIST → PENDING → ACCEPTED / DECLINED / REMOVED",    "send, accept, decline, unfriend"],
        ["authSlice (Redux)",            "SM-06", "UNAUTHENTICATED → LOADING → AUTHENTICATED / ERROR",            "loginUser, logout, 401 interceptor"],
        ["ActivityPage (React)",         "SM-07", "LOADING → CONTENT_READY → ANSWERING → SUBMITTING → PASSED/FAILED", "content fetch, submit, retry"],
        ["WhisperService (Python)",      "SM-08", "UNINITIALIZED → LOADING → AVAILABLE / UNAVAILABLE → TRANSCRIBING", "first transcribe_audio() call"],
        ["GroqService call (Python)",    "SM-09", "IDLE → BUILDING_PROMPT → CALLING_GROQ → PARSING → RESULT/ERROR", "validate_activity(), generate_tier_feedback()"],
        ["AudioRecorder (React)",        "SM-10", "READY → RECORDING → PREVIEW → UPLOADING → TRANSCRIBED",       "record, stop, submit, re-record"],
        ["ProtectedRoute / AdminRoute",  "SM-11", "CHECKING → AUTHORIZED / UNAUTHORIZED",                         "route change, auth state change"],
        ["progressSlice (Redux)",        "SM-12", "EMPTY → LOADING_PROGRESS → LOADED",                           "fetchAllProgress, startPair, advanceProgress"],
        ["AdminDashboard (React)",       "SM-13", "LOADING → VIEWING → EDITING (per section) → SAVING",           "page load, item select, save, delete"],
    ],
    header_color="1F3864"
)

doc.add_paragraph()
add_body("End of LearnWise Software Design Document v1.0", bold=True, italic=True)
add_body("Prepared for Stakeholder & Client Presentation — April 2026", italic=True)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
output_path = "/home/jeel/Downloads/learnwise-react_till_now_6/learnwise-react/docs/DESIGN_DOCUMENT.docx"
doc.save(output_path)
print(f"SUCCESS: Document saved to {output_path}")
